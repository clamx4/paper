'''
Created on 2015年12月28日

@author: cdz
'''

from http import cookies
import urllib
from urllib.request import urlopen, Request
from urllib.parse import urlencode
from http.client import HTTPConnection
from time import sleep
import urllib.parse
from bs4 import BeautifulSoup
from docx.api import Document
import re



def getQueryHtml(keyword):
    global COOKIES
    location = getWhileTrue(getLocation, keyword)
    req = Request(location)
    req.add_header('Cookie', COOKIES.output(header='', sep=';'))
    page = urlopen(req, timeout=10)
    html = str(page.read(), encoding='utf-8')
    return html

def getDetailPageHTML(url):
    global COOKIES
    host = 'apps.webofknowledge.com'
    headers = {'Accept':'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
               'Accept-Encoding':'gzip, deflate, sdch',
               'Accept-Language':'zh-CN,zh;q=0.8',
               'Connection':'keep-alive',
               'Cookie':COOKIES.output(header='', sep=';'),
               'Host':'apps.webofknowledge.com',
               'User-Agent':'Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/39.0.2171.99 Safari/537.36'
               }
    conn = HTTPConnection(host)
    conn.request('GET', url, None, headers)
    resp = conn.getresponse()
    headers = resp.headers
    if resp.status == 302:
        location = headers.get('Location')
        for cookie in headers.get_all('Set-Cookie'):
            COOKIES.load(cookie)
        conn = HTTPConnection('apps.webofknowledge.com')
        conn.request('GET', location, None, {'Cookie' : COOKIES.output(header='', sep=';')})
        resp = conn.getresponse()
        headers = resp.headers
        
    for cookie in headers.get_all('Set-Cookie'):
        COOKIES.load(cookie)
    html = resp.read()
    return str(html, encoding='utf-8')

def getWhileTrue(func, *args, interval=10):
    while True:
        try:
            result = func(*args)
            break
        except Exception as e:
            sleep(interval)
            print('error when', str(func), ':', e)
    return result

def getSID():
    '''store SID and cookies in global var COOKIES'''
    global COOKIES
    host = 'www.webofknowledge.com'
    headers = {'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
               'Referer': 'http://nav.lib.xjtu.edu.cn/database/index.do?action=redirect&pid=198',
                'Accept-Encoding': 'gzip, deflate, sdch',
                'Accept-Language': 'zh-CN,zh;q=0.8',
                'Connection': 'keep-alive',
                'Host': 'www.webofknowledge.com',
                'User-Agent': 'Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/39.0.2171.99 Safari/537.36',
               }
    conn = HTTPConnection(host)
    conn.request('GET', '', None, headers)
    resp = conn.getresponse()
    
    # redirect to http://apps.webofknowledge.com/?SID=XXXXX&SrcApp=CR&Init=Yes
    headers = resp.headers
    location = headers.get('Location')
    #cookies = headers.get_all('Set-Cookie')
    for cookie in headers.get_all('Set-Cookie'):
        COOKIES.load(cookie)
    conn = HTTPConnection('apps.webofknowledge.com')
    conn.request('GET', location, None, {'Cookie' : COOKIES.output(header='', sep=';')})
    resp = conn.getresponse()
    
    # redirect to http://apps.webofknowledge.com/home.do;jsessionid=XXXX?SID=XXXX&SrcApp=CR&Init=Yes
    headers = resp.headers
    location = headers.get('Location')
    #cookies_jsessionid = headers.get_all('Set-Cookie')
    for cookie in headers.get_all('Set-Cookie'):
        COOKIES.load(cookie)
    conn = HTTPConnection('apps.webofknowledge.com')
    conn.request('GET', location, None, {'Cookie' : COOKIES.output(header='', sep=';')})
    resp = conn.getresponse()
    
    # redirect to http://apps.webofknowledge.com/UA_GeneralSearch_input.do?product=UA&search_mode=GeneralSearch&SID=XXX&preferencesSaved= 
    headers = resp.headers
    location = headers.get('Location')
    #cookies_jsessionid = headers.get_all('Set-Cookie')
    for cookie in headers.get_all('Set-Cookie'):
        COOKIES.load(cookie)
    conn = HTTPConnection('apps.webofknowledge.com')
    conn.request('GET', location, None, {'Cookie' : COOKIES.output(header='', sep=';')})
    resp = conn.getresponse()
    
    headers = resp.headers
    location = headers.get('Location')
    for cookie in headers.get_all('Set-Cookie'):
        COOKIES.load(cookie)

def getLocation(keyword):
    global COOKIES
    sid = COOKIES['SID'].value
    data = {'fieldCount':'1',
            'action':'search',
            'product':'UA',
            'search_mode':'GeneralSearch',
            'SID':sid,
            'max_field_count':'25',
            'max_field_notice':'注意: 无法添加另一字段。',
            'input_invalid_notice':'检索错误: 请输入检索词。',
            'exp_notice':'检索错误: 专利检索词可在多个家族中找到 (',
            'input_invalid_notice_limits':' <br/>注: 滚动框中显示的字段必须至少与一个其他检索字段相组配。',
            'sa_params':'UA||' + sid + '|http://apps.webofknowledge.com|\'',
            'formUpdated':'true',
            'value(input1)':keyword,
            'value(select1)':'TI',
            'x':'718',
            'y':'294',
            'value(hidInput1)':'',
            'limitStatus':'collapsed',
            'ss_lemmatization':'On',
            'ss_spellchecking':'Suggest',
            'SinceLastVisit_UTC':'',
            'SinceLastVisit_DATE':'',
            'period':'Range Selection',
            'range':'ALL',
            'startYear':'1900',
            'endYear':'2015',
            'update_back2search_link_param':'yes',
            'ssStatus':'display:none',
            'ss_showsuggestions':'ON',
            'ss_query_language':'auto',
            'ss_numDefaultGeneralSearchFields':'1',
            'rs_sort_by':'PY.D;LD.D;SO.A;VL.D;PG.A;AU.A'
            }
    url = 'http://apps.webofknowledge.com/UA_GeneralSearch.do'
    host = 'apps.webofknowledge.com'
    conn = HTTPConnection(host)
    headers = {'Host':'apps.webofknowledge.com',
               'Connection':'keep-alive',
               'Cache-Control': 'max-age=0',
               'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
               'Origin': 'http://apps.webofknowledge.com',
               'User-Agent': 'Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/39.0.2171.99 Safari/537.36',
               'Content-Type': 'application/x-www-form-urlencoded',
               'Referer': 'http://apps.webofknowledge.com/UA_GeneralSearch_input.do?product=UA&search_mode=GeneralSearch&SID=' + sid + '&preferencesSaved=',
               'Accept-Encoding': 'gzip, deflate',
               'Accept-Language': 'zh-CN,zh;q=0.8',
               'Cookie':COOKIES.output(header='', sep=';')
               }
    conn.request('POST', url, urlencode(data).encode(encoding='utf_8', errors='strict'), headers)
    resp = conn.getresponse()
    headers = resp.headers
    for cookie in headers.get_all('Set-Cookie'):
        COOKIES.load(cookie)
    return resp.headers.get('Location')

def getLinksFromQueryHTML(queryHtml):
    bs = BeautifulSoup(queryHtml, 'html.parser')
    host = 'http://apps.webofknowledge.com'
    result = []
    for i in range(1, 11):
        _id = 'RECORD_' + str(i)
        record = bs.find(id=_id)
        recordSoup = BeautifulSoup(str(record), 'html.parser')
        alinks = recordSoup.find_all('a', {'class' : 'smallV110'})
        if len(alinks) == 1:
            link = host + alinks[0]['href']
            result.append(link)
        else:
            raise Exception()
        sleep(10)
    return result

def doRecord(doc, detailHTML, num):
    bs = BeautifulSoup(detailHTML, 'html.parser')
    title = bs.find_all('div', {'class' : 'title'})
    if len(title) == 1:
        title_content = re.sub(r'<.*?>', '', str(title[0].contents[1]))
        #print(title_content)
        doc.add_heading(str(num) + '. ' + title_content, 1)
        author = title[0].next_sibling
        nosup = re.sub(r'<sup>.*?</sup>', '', str(author).replace('\n', ''))
        author_content = re.sub(r'<.*?>', '', nosup)
        author_content = re.sub(r'更多内容.*', '', author_content)
        #print(author_content)
        doc.add_paragraph(author_content)
        info_fusion = author.next_sibling
        info_fusion_content = re.sub(r'[\n]+', '\n', re.sub(r'<.*?>', '', str(info_fusion)))
        info_fusion_content = info_fusion_content.replace('卷:\n', '卷:').replace('\n期:\n', '期:').replace('\n页:\n', '页:').replace('子辑:\n', '子辑:')
        info_fusion_content = info_fusion_content.replace('DOI:\n', 'DOI:')
        info_fusion_content = info_fusion_content.replace('出版年:\n', '出版年:')
        info_fusion_content = info_fusion_content.replace('ISSN: \n', 'ISSN: ')
        info_fusion_content = info_fusion_content.replace('eISSN: \n', 'eISSN: ')
        info_fusion_content = info_fusion_content.replace('&amp;', '&')
        #print(info_fusion_content)
        doc.add_paragraph(info_fusion_content)
        abstract = info_fusion.next_sibling
        abstract_content = re.sub(r'[\n]+', '\n', re.sub(r'<.*?>', '', str(abstract)))
        #print(abstract_content)
        doc.add_paragraph(abstract_content)
        keyword = abstract.next_sibling
        keyword_content = re.sub(r'[\n]+', '\n', re.sub(r'<.*?>', '', str(keyword)))
        #print(keyword_content)
        doc.add_paragraph(keyword_content)
    elif len(title) == 0:
        #title = bs.find_all('table', {'id' : 'FullRecDataTable'})
        #title = bs.find(id='FullRecDataTable')
        title = bs.find_all('td', {'class' : 'FullRecTitle'})
        title_content = re.sub(r'<.*?>', '', str(title[0]))
        info_rows = bs.find_all('td', {'class' : 'fr_data_row'})
        for info in info_rows:
            info_content = ''.join([str for str in info.stripped_strings])
            if 'Patent Number(s):' in info_content:
                专利号 = info_content.replace('Patent Number(s):', '专利号:')
            elif 'Inventor(s):' in info_content:
                发明人 = info_content.replace('Inventor(s):', '发明人:')
            elif 'Patent Assignee Name(s) and Code(s):' in info_content:
                专利权人和代码 = info_content.replace('Patent Assignee Name(s) and Code(s):', '专利权人和代码:')
            elif 'Derwent Primary Accession Number:' in info_content:
                Derwent主入藏号 = info_content.replace('Derwent Primary Accession Number:', 'Derwent主入藏号:')
            elif 'Abstract:' in info_content:
                摘要 = info_content.replace('Abstract:', '摘要:\n')
            elif 'International Patent Classification:' in info_content:
                国际专利分类 = info_content.replace('International Patent Classification:', '国际专利分类:')
            elif 'Derwent Class Code(s):' in info_content:
                德温特分类代码 = info_content.replace('Derwent Class Code(s):', '德温特分类代码:')
            elif 'Derwent Manual Code(s):' in info_content:
                德温特手工代码 = info_content.replace('Derwent Manual Code(s):', '德温特手工代码:')
        doc.add_heading(str(num) + '. ' + title_content, 1)
        if 专利号:
            doc.add_paragraph(专利号)
        if 发明人:
            doc.add_paragraph(发明人)
        if 专利权人和代码:
            doc.add_paragraph(专利权人和代码)
        if Derwent主入藏号:
            doc.add_paragraph(Derwent主入藏号)
        if 摘要:
            doc.add_paragraph(摘要)
        if 国际专利分类:
            doc.add_paragraph(国际专利分类)
        if 德温特分类代码:
            doc.add_paragraph(德温特分类代码)
        if 德温特手工代码:
            doc.add_paragraph(德温特手工代码)
            
    
COOKIES = cookies.SimpleCookie()
def start(keyword, paper_from, paper_to, store_path='d:/论文摘要信息-wok.docx'):
    global COOKIES
    getWhileTrue(getSID)
    queryHtml = getWhileTrue(getQueryHtml, keyword)
        
    doc = Document()
    doc.add_heading('关键词：' + keyword, 0)
    doc.save(store_path)
    for i in range(paper_from, paper_to + 1):
        print('getting', i , '/', paper_to, 'paper')
        sid = COOKIES['SID'].value
        link = ''.join(['http://apps.webofknowledge.com/full_record.do?',
                'product=UA',
                '&search_mode=GeneralSearch',
                '&qid=1',
                '&SID=',sid,
                '&page=1',
                '&doc=', str(i)
                ])
        try:
            detailHTML = getWhileTrue(getDetailPageHTML, link)
            sleep(10)
        except Exception as e:
            print('error when getting detail paper infomation:', e)
            getWhileTrue(getSID)
            getWhileTrue(getQueryHtml, keyword)
            sleep(15)
        doc = Document(store_path)
        doRecord(doc, detailHTML, i)
        doc.save(store_path)
        
if __name__ == '__main__':
    #start(keyword, paper_from, paper_to)
    pass