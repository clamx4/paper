'''
Created on 2015年12月1日

@author: cdz
'''

from PIL import Image
import time
from bs4 import BeautifulSoup
from os import path
from urllib.request import urlopen, Request
from http import cookies
from docx import Document
import re


def getHtml(keyword, page=1):
    url = 'http://dl.acm.org/results.cfm?query=' + keyword + '&start=' + str(20 * (page - 1)) + '&filtered=&within=owners%2Eowner%3DACM&dte=&bfr=&srt=_score'
    req = Request(url)
    req.add_header('User-Agent', 'Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/39.0.2171.99 Safari/537.36')
    
    while True:
        try:
            page = urlopen(req)
            break
        except Exception as e:
            print(e)
    html = page.read()
    return html

def getTitleAndPath(thediv):
    divsoup = BeautifulSoup(str(thediv), 'html.parser')
    titlediv = divsoup.find('div', 'title')
    title = titlediv.a.string
    print(title)
    alist = divsoup.find_all('a')
    pdfLinkList = [a for a in alist if 'name' in a.attrs.keys()]
    if len(pdfLinkList) < 1:
        path = None
    else:
        path = pdfLinkList[0].attrs['href']
    print(path)
    return title, path

def download(title, path, folder):
    print('try to get', title, path)
    host = 'http://dl.acm.org/'
    url = host + path
    req = Request(url)
    req.add_header('User-Agent', 'Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/39.0.2171.99 Safari/537.36')
    while True:
        try:
            page = urlopen(req, timeout=8)
            pdfbytes = page.read()
            break
        except Exception as e:
            print(e)
    filename = getFileName(title)
    f = open(folder + filename, 'wb')
    f.write(pdfbytes)
    f.close()
    
def getFileName(title):
    file_name = title + '.pdf'
    file_name = file_name.replace('\\', '')
    file_name = file_name.replace('/', '')
    file_name = file_name.replace(':', '')
    file_name = file_name.replace('*', '')
    file_name = file_name.replace('?', '')
    file_name = file_name.replace('\"', '')
    file_name = file_name.replace('<', '')
    file_name = file_name.replace('>', '')
    file_name = file_name.replace('|', '')
    return file_name

def start(keyword, page_from, page_to, folder):
    for page in range(page_from, page_to + 1):
        html = getHtml(keyword, page)
        soup = BeautifulSoup(html, 'html.parser')
        divList = soup.find_all('div', {'class' : 'details'})
        title2path = {}
        for thediv in divList:
            title, path = getTitleAndPath(thediv)
            if path != None:
                title2path[title] = path
        papernum = len(title2path)
        i = 1
        for title, path in title2path.items():
            print('downloading paper' , str(i), '/', str(papernum), 'in page', str(page), '/', str(page_to))
            download(title, path, folder)
            time.sleep(10)
            i += 1
            
    
def dropTail(string):
    arr = string.replace('try to get ', '').split(' ')
    title = ' '.join(arr[0 : len(arr) - 1])
    return getFileName(title)

def exsit(filename):
    folder = 'd:/p/'
    filepath = folder + filename;
    if path.isfile(filepath):
        print('exsit')
    else:
        print('x')
        
def func(arg1, arg2, arg3):
    print(arg1, arg2, arg3)
    
def proxy(foo):
    def newfunc(*args):
        return foo(1, *args)
    return newfunc

def ret():
    return 1, 2

def foo(q,w):
    print(q)
    print(w)

def main():
    a = [x / 2 for x in [1,2,3]]
    print(a)
    pass

def test(id='1'):
    print(id)
    
def doRecord(doc, detailHTML, num):
#     f = open('d:/a1.htm', encoding='utf-8')
#     detailHTML = f.read()
#     f.close()
    bs = BeautifulSoup(detailHTML, 'html.parser')
    title = bs.find_all('div', {'class' : 'title'})
    title_content = re.sub(r'<.*?>', '', str(title[0].contents[1]))
    print(title_content)
    #doc.add_paragraph(title_content)
    doc.add_heading(str(num) + '. ' + title_content, 1)
    author = title[0].next_sibling
    nosup = re.sub(r'<sup>.*?</sup>', '', str(author).replace('\n', ''))
    print(re.sub(r'<.*?>', '', nosup))
    doc.add_paragraph(re.sub(r'<.*?>', '', nosup))
    info_fusion = author.next_sibling
    info_fusion_content = re.sub(r'[\n]+', '\n', re.sub(r'<.*?>', '', str(info_fusion)))
    info_fusion_content = info_fusion_content.replace('卷:\n', '卷:').replace('\n期:\n', '期:').replace('\n页:\n', '页:')
    info_fusion_content = info_fusion_content.replace('DOI:\n', 'DOI:')
    info_fusion_content = info_fusion_content.replace('出版年:\n', '出版年:')
    info_fusion_content = info_fusion_content.replace('ISSN: \n', 'ISSN: ')
    info_fusion_content = info_fusion_content.replace('eISSN: \n', 'eISSN: ')
    print(info_fusion_content)
    doc.add_paragraph(info_fusion_content)
    abstract = info_fusion.next_sibling
    abstract_content = re.sub(r'[\n]+', '\n', re.sub(r'<.*?>', '', str(abstract)))
    print(abstract_content)
    doc.add_paragraph(abstract_content)
    keyword = abstract.next_sibling
    keyword_content = re.sub(r'[\n]+', '\n', re.sub(r'<.*?>', '', str(keyword)))
    print(keyword_content)
    doc.add_paragraph(keyword_content)



if __name__ == '__main__':
#     doc = Document()
#     doc.save('d:/x.docx')
#     doc = Document('d:/x.docx')
#     doRecord(doc)
#     doc.save('d:/x.docx')
    s = '123\n3'
    print(re.sub(r'2.*', '', s))
    


